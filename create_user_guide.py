import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_guide():
    doc = docx.Document()

    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Base Styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B) # Slate-800

    # Colors
    ORANGE_COLOR = RGBColor(0xF9, 0x73, 0x16) # Orange-500
    NAVY_COLOR = RGBColor(0x0F, 0x17, 0x2A) # Slate-900
    GOLD_COLOR = RGBColor(0xD9, 0x77, 0x06) # Amber-600

    # ── DOCUMENT HEADER TITLE ─────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("Oogway AI — User Guide & Session Analysis Manual")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY_COLOR

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = subtitle_p.add_run("Comprehensive Operational Guide for Analyzing Live Class Intelligence at Kraftshala")
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = GOLD_COLOR

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── EXECUTIVE SUMMARY BOX ─────────────────────────────────────────────
    table_exec = doc.add_table(rows=1, cols=1)
    table_exec.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_exec = table_exec.cell(0, 0)
    set_cell_background(cell_exec, "FFF7ED") # Amber-50
    set_cell_margins(cell_exec, top=140, bottom=140, left=180, right=180)

    p_exec = cell_exec.paragraphs[0]
    p_exec.paragraph_format.space_after = Pt(4)
    r_exec_head = p_exec.add_run("📌 Executive Overview: What is Oogway AI?")
    r_exec_head.bold = True
    r_exec_head.font.size = Pt(12)
    r_exec_head.font.color.rgb = ORANGE_COLOR

    p_exec_body = cell_exec.add_paragraph()
    p_exec_body.paragraph_format.space_after = Pt(2)
    p_exec_body.add_run(
        "Oogway AI is Kraftshala's proprietary session intelligence platform. It ingests raw live class video recordings "
        "and transcripts (.vtt) and automatically executes multi-stage AI audits to evaluate trainer quality, "
        "identify student confusion points, and track batch mastery levels across 6 key quality dimensions."
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── SECTION 1: CORE ARCHITECTURE & WORKFLOW ─────────────────────────────
    h1 = doc.add_heading("1. Oogway Analysis Architecture & Flow", level=1)
    h1.runs[0].font.color.rgb = NAVY_COLOR
    h1.runs[0].font.size = Pt(16)

    doc.add_paragraph(
        "When a session is uploaded or ingested into Oogway, the platform runs a 4-tiered analysis pipeline:"
    )

    # Summary Table of Tabs
    table_tabs = doc.add_table(rows=5, cols=3)
    table_tabs.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tabs.autofit = False

    headers = ["Tab Name", "Primary Purpose", "Key Metrics & Outputs"]
    hdr_cells = table_tabs.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "0F172A")
        p = hdr_cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data_rows = [
        ("⏱ Timeline", "Chronological Milestone Mapping", "Major session milestones, administrative delay flags, timestamp navigation."),
        ("⚡ Oogway Pulse", "360° Executive Summary", "Instant high-level verdict, top expert strengths, student engagement rating."),
        ("🎯 Expert Go", "Deep Expert Quality Audit", "6 SST Dimensions, Kraftshala Quality Checklist, Single Synthesized Cards, Feedback Emails."),
        ("🎓 Student Go", "Cohort Doubt & Mastery Audit", "Batch Mastery Level, Timestamped Confusion Points, Question Classification, Next Steps."),
    ]

    for row_idx, data in enumerate(data_rows, start=1):
        row_cells = table_tabs.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F8FAFC")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ── SECTION 2: DETAILED FEATURE GUIDE & HOW TO ANALYZE ─────────────────
    h2 = doc.add_heading("2. Step-by-Step Data Analysis Guide", level=1)
    h2.runs[0].font.color.rgb = NAVY_COLOR
    h2.runs[0].font.size = Pt(16)

    # Feature 1: Timeline
    h2_1 = doc.add_heading("Feature 1: Session Flow Timeline (Pacing & Milestones)", level=2)
    h2_1.runs[0].font.color.rgb = ORANGE_COLOR

    p1 = doc.add_paragraph()
    p1.add_run("• What it does: ").bold = True
    p1.add_run("Parses the transcript chronologically into distinct learning segments (e.g. Session Start, Concept Explanation, Q&A, Admin Break).\n")
    p1.add_run("• How to use it: ").bold = True
    p1.add_run("Look for red 'ISSUE FLAGGED' tags indicating administrative delays or topic dragging. Click any timestamp range (e.g. 00:03:52 ➔ 00:14:37) to inspect the recording.")

    # Feature 2: Oogway Pulse
    h2_2 = doc.add_heading("Feature 2: Oogway Pulse (Executive Synthesis)", level=2)
    h2_2.runs[0].font.color.rgb = ORANGE_COLOR

    p2 = doc.add_paragraph()
    p2.add_run("• What it does: ").bold = True
    p2.add_run("Provides an instant 60-second summary of overall session health.\n")
    p2.add_run("• Key Highlights: ").bold = True
    p2.add_run("Camera/Punctuality hygiene badges, core topic summary, expert strengths, and student engagement index.")

    # Feature 3: Expert Go
    h2_3 = doc.add_heading("Feature 3: Expert Go (6-Dimension Quality Audit)", level=2)
    h2_3.runs[0].font.color.rgb = ORANGE_COLOR

    p3 = doc.add_paragraph()
    p3.add_run("• Evaluated Dimensions:\n").bold = True
    p3.add_run(
        "  1. Content Accuracy & Depth (Factual correctness, industry terminology)\n"
        "  2. Pedagogical Approach (Structure, active learning, analogies used)\n"
        "  3. Live Platform Walkthrough (UI screen navigation, hands-on demos)\n"
        "  4. Pacing & Time Management (Class speed, delay handling)\n"
        "  5. Student Emotional Support (Patience, encouragement, psychological safety)\n"
        "  6. Delivery Fluency (Clarity of speech, language, filler words)\n\n"
    )
    p3.add_run("• Single Synthesized Dimension Card: ").bold = True
    p3.add_run("Under each dimension, all observations are unified into ONE structured card containing What Happened, Why It Matters, and Action Recommendations for Next Session.\n")
    p3.add_run("• Automated Feedback Email Generator: ").bold = True
    p3.add_run("Click 'Send Feedback Email ✉️' to open a pre-filled Gmail draft available in Variant A (Warm Coaching) or Variant B (Direct Actionable).")

    # Feature 4: Student Go
    h2_4 = doc.add_heading("Feature 4: Student Go (Cohort Mastery & Confusion Audit)", level=2)
    h2_4.runs[0].font.color.rgb = ORANGE_COLOR

    p4 = doc.add_paragraph()
    p4.add_run("• Batch Mastery Level: ").bold = True
    p4.add_run("Classified as Advanced, Competent, Developing, or Needs Foundations.\n")
    p4.add_run("• Timestamped Confusion Points: ").bold = True
    p4.add_run("Tracks exact moments where students expressed confusion, asked for repetition, or paused the expert.\n")
    p4.add_run("• Question Classification: ").bold = True
    p4.add_run("Categorizes all student questions into Conceptual, Operational, Live Platform, and Clarification doubts.")

    # Feature 5: Source Material & AI Chat
    h2_5 = doc.add_heading("Feature 5: Source Material & Ask Master Oogway AI", level=2)
    h2_5.runs[0].font.color.rgb = ORANGE_COLOR

    p5 = doc.add_paragraph()
    p5.add_run("• View Source Drawer: ").bold = True
    p5.add_run("Includes the embedded video player with seek functionality and copyable/downloadable raw transcript (.vtt).\n")
    p5.add_run("• Ask Master Oogway: ").bold = True
    p5.add_run("Floating chatbot widget to ask any custom question about the class (e.g., 'Did the expert explain CBO vs ABO correctly?').")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # ── SECTION 3: OPERATIONAL CHECKLIST FOR PROGRAM LEADS ─────────────────
    h3 = doc.add_heading("3. Program Lead & Mentor Operating Checklist", level=1)
    h3.runs[0].font.color.rgb = NAVY_COLOR
    h3.runs[0].font.size = Pt(16)

    table_check = doc.add_table(rows=5, cols=3)
    table_check.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_check.autofit = False

    ch_headers = ["Step", "Action Item", "Target Metric / Tool"]
    for i, title in enumerate(ch_headers):
        table_check.rows[0].cells[i].text = title
        set_cell_background(table_check.rows[0].cells[i], "0F172A")
        p = table_check.rows[0].cells[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    check_data = [
        ("Step 1", "Review Session Flow & Delay Flags", "Timeline Tab -> Check 'ISSUE FLAGGED'"),
        ("Step 2", "Audit Expert Teaching Gaps", "Expert Go Tab -> Check Severity Badges & Action Recommendations"),
        ("Step 3", "Review Student Confusion Friction", "Student Go Tab -> Timestamped Confusion Points"),
        ("Step 4", "Send Feedback Email to Trainer", "Expert Go Tab -> Click 'Send Feedback Email ✉️'"),
    ]

    for row_idx, data in enumerate(check_data, start=1):
        row_cells = table_check.rows[row_idx].cells
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            if row_idx % 2 == 0:
                set_cell_background(row_cells[col_idx], "F8FAFC")

    doc.save("Oogway_User_Guide_and_Analysis_Manual.docx")
    print("Oogway_User_Guide_and_Analysis_Manual.docx generated successfully!")

if __name__ == '__main__':
    create_guide()
